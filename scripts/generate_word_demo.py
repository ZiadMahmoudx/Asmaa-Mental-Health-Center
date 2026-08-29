import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def create_demo_doc(output_path):
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    TEAL_DARK = RGBColor(19, 78, 74)     # #134e4a
    TEAL_MAIN = RGBColor(13, 148, 136)   # #0d9488
    SLATE_DARK = RGBColor(30, 41, 59)    # #1e293b
    GRAY_TEXT = RGBColor(71, 85, 105)    # #475569

    # Title
    p_title = doc.add_paragraph()
    set_rtl(p_title)
    r_title = p_title.add_run("مركز أسما للصحة النفسية — Asmaa Mental Health Clinic")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = TEAL_DARK

    p_sub = doc.add_paragraph()
    set_rtl(p_sub)
    r_sub = p_sub.add_run("دليل السيناريو العملي لتجربة المنصة واستعراض المزايا (Live Demo Guide)")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.bold = True
    r_sub.font.color.rgb = TEAL_MAIN

    p_intro = doc.add_paragraph()
    set_rtl(p_intro)
    r_intro = p_intro.add_run("هذا الدليل مخصص لاستعراض النظام المباشر خطوة بخطوة للعميل، ويغطي دورة العمل الكاملة بدءاً من تقييم وحجز المريض، مروراً بمراجعة الإدارة والمدفوعات، وصولاً إلى الملف السريري والتقارير الطبية للاستشاري.")
    r_intro.font.name = "Arial"
    r_intro.font.size = Pt(10.5)
    r_intro.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    # SECTION 1: CREDENTIALS
    h1 = doc.add_paragraph()
    set_rtl(h1)
    r_h1 = h1.add_run("🔑 أولاً: بيانات الدخول التجريبية الجاهزة")
    r_h1.font.name = "Arial"
    r_h1.font.size = Pt(14)
    r_h1.font.bold = True
    r_h1.font.color.rgb = TEAL_DARK

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["الدور والصلاحية", "البريد الإلكتروني", "كلمة المرور"]
    col_widths = [Inches(2.0), Inches(3.0), Inches(1.8)]

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "0F766E")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        set_rtl(p)
        run = p.add_run(title)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("لوحة الإدارة والاستقبال (Admin)", "admin@asmaaclinic.com", "AsmaaAdmin2026"),
        ("لوحة الاستشاري الطبي (Doctor)", "dr.asmaa@asmaaclinic.com", "AsmaaDoctor2026"),
        ("بوابة المريض (Patient)", "sara.mahmoud@example.com", "AsmaaPatient2026"),
    ]

    for row_idx, row_data in enumerate(data, start=1):
        cells = table.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cells[col_idx].width = col_widths[col_idx]
            set_cell_background(cells[col_idx], bg_color)
            set_cell_margins(cells[col_idx], top=100, bottom=100, left=140, right=140)
            p = cells[col_idx].paragraphs[0]
            set_rtl(p)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = SLATE_DARK
            else:
                run.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph()

    # SECTION 2: DEMO SCENARIO
    h2 = doc.add_paragraph()
    set_rtl(h2)
    r_h2 = h2.add_run("🎬 ثانياً: السيناريو العملي المتسلسل للعرض المباشر")
    r_h2.font.name = "Arial"
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = TEAL_DARK

    # Phase 1
    p_p1 = doc.add_paragraph()
    set_rtl(p_p1)
    r_p1 = p_p1.add_run("1. رحلة المريض (Patient Journey):")
    r_p1.font.name = "Arial"
    r_p1.font.size = Pt(12)
    r_p1.font.bold = True
    r_p1.font.color.rgb = TEAL_MAIN

    steps_p1 = [
        ("التقييم النفسي المعتمد (/assessments): ", "يقوم المريض بملء مقياس الاكتئاب السريري (PHQ-9) أو القلق (GAD-7) أو مؤشر الأرق (ISI). يقوم النظام بحساب الدرجة تلقائياً وتحديد مستوى الشدة ورصد إشارات الخطر على السيرفر بدقة إكلينيكية دون الاعتماد على مدخلات المتصفح."),
        ("حجز الجلسة واختيار الموعد (/booking): ", "اختيار الطبيب المناسب (د. أسماء عبد الوهاب)، ونوع الجلسة (أونلاين عبر زووم أو حضورية بالعيادة)، واختيار اليوم والساعة المتاحة بتوقيت القاهرة المحلي."),
        ("حجز المقعد وتفاصيل الدفع (/payment): ", "حجز فوري للوقت ومنع أي مريض آخر من حجزه (Anti-Double Booking)، وعرض بيانات الدفع (InstaPay / فودافون كاش) مع مؤقت عد تنازلي لحفظ الحجز."),
        ("رفع إيصال التحويل: ", "رفع صورة إيصال التحويل مع فحص أمان للملف للتحقق من سلامته قبل حفظه."),
    ]

    for title, desc in steps_p1:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"• {title}")
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = SLATE_DARK
        r2 = p.add_run(desc)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    # Phase 2
    p_p2 = doc.add_paragraph()
    set_rtl(p_p2)
    r_p2 = p_p2.add_run("2. مكتب الإدارة والاستقبال (Admin & Operations):")
    r_p2.font.name = "Arial"
    r_p2.font.size = Pt(12)
    r_p2.font.bold = True
    r_p2.font.color.rgb = TEAL_MAIN

    steps_p2 = [
        ("لوحة مؤشرات المركز (/dashboard/admin): ", "عرض الإحصائيات الحقيقية المحدثة لحظياً (الجلسات المكتملة، الإيرادات، خط الإيراد المؤكد، حالات الفرز العاجلة، ونسبة جلسات الأونلاين مقارنة بالعيادة)."),
        ("مكتب تدقيق المدفوعات (/dashboard/admin/verification): ", "مراجعة إيصال المريض، مطابقة رقم العملية، والضغط على 'تأكيد الحجز' مع إمكانية فتح محادثة واتساب فورية مجهزة بنص التأكيد والموعد بتوقيت القاهرة."),
        ("إدارة جداول ومواعيد الأطباء (/dashboard/admin/schedule): ", "إمكانية قيام الإدارة بالتحكم الكامل في فترات العمل الأسبوعية لأي طبيب، إضافة أو تعديل أو إلغاء فترات، وتسجيل إجازات العيادة والإغلاق المؤقت مع فحص التضارب التلقائي."),
        ("سجل العمليات والتحكم في الحجوزات (/dashboard/admin/appointments): ", "شاشة بحث وتصفية شاملة لجميع الحجوزات، مع إمكانية إعادة الجدولة أو الإلغاء أو إضافة روابط زووم أو تحرير الحجوزات المعلقة."),
    ]

    for title, desc in steps_p2:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"• {title}")
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = SLATE_DARK
        r2 = p.add_run(desc)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    # Phase 3
    p_p3 = doc.add_paragraph()
    set_rtl(p_p3)
    r_p3 = p_p3.add_run("3. مساحة عمل الاستشاري والملف السريري (Doctor Workspace):")
    r_p3.font.name = "Arial"
    r_p3.font.size = Pt(12)
    r_p3.font.bold = True
    r_p3.font.color.rgb = TEAL_MAIN

    steps_p3 = [
        ("الجدول اليومي وقائمة الجلسات: ", "استعراض جلسات اليوم القادمة والمكتملة، مع رابط مباشر للجلسة عبر زووم، وزر تذكير المريض عبر الواتساب."),
        ("الملف السريري الذكي للمريض (Clinical Drawer): ", "بضغطة واحدة يفتح درج جانبي يعرض مسار درجات المقاييس النفسية السابقة (PHQ-9, GAD-7) كرسوم بيانية، مع عرض خطة الأمان (Safety Plan) وأرقام التواصل الموثوقة، وتاريخ تشخيصات الجلسات السابقة."),
        ("توثيق وتوقيع التقرير الطبي القانوني (SOAP Note): ", "كتابة الشكوى، التشخيص، أكواد DSM-5 المعتمدة، الخطة الدوائية، وتوصيات المتابعة، مع إمكانية التوقيع الإلكتروني النهائي غير القابل للتعديل لضمان الموثوقية القانونية."),
        ("إدارة مواعيد العمل والإجازات: ", "سهولة إضافة أو تعديل نوافذ العمل الأسبوعية، تسجيل فترات السفر والمؤتمرات، وإعادة جدولة أي موعد مع إرسال إشعار واتساب تلقائي بالموعدين القديم والجديد للمريض."),
    ]

    for title, desc in steps_p3:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"• {title}")
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = SLATE_DARK
        r2 = p.add_run(desc)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    # SECTION 3: KEY DIFFERENTIATORS
    h3 = doc.add_paragraph()
    set_rtl(h3)
    r_h3 = h3.add_run("🌟 ثالثاً: نقاط القوة والتميز في النظام الطبي للمركز")
    r_h3.font.name = "Arial"
    r_h3.font.size = Pt(14)
    r_h3.font.bold = True
    r_h3.font.color.rgb = TEAL_DARK

    points = [
        ("🔒 حماية التضارب على مستوى قاعدة البيانات: ", "استحالة حدوث حجز مزدوج لنفس الدقيقة بين مريضين عبر قفل الأقفال المركب (Composite Lock Key)."),
        ("⏰ دقة التوقيت وثباته: ", "تخزين كافة المواعيد بنظام UTC العالمي وعرضها بتوقيت القاهرة المحلي تلقائياً، لمنع أي انزياح في المواعيد عند تطبيق التوقيت الصيفي."),
        ("📱 تكامل الواتساب الفوري: ", "رسائل جاهزة بنقرة واحدة لتأكيد الحجز، التذكير بالموعد، إرسال رابط زووم، أو إشعار إعادة الجدولة."),
        ("📋 خطة الأمان السريرية (Stanley-Brown): ", "أداة تفاعلية لحماية المرضى المعرضين للخطر، متاحة للمريض والاستشاري المعالج في أي لحظة حرجة."),
    ]

    for title, desc in points:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"✔ {title}")
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = TEAL_DARK
        r2 = p.add_run(desc)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY_TEXT

    # Footer note
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    set_rtl(p_foot)
    r_foot = p_foot.add_run("تم إعداد هذا المستند كدليل تشغيلي متكامل لمنصة مركز أسما للصحة النفسية © 2026")
    r_foot.font.name = "Arial"
    r_foot.font.size = Pt(9)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(output_path)
    print("DOCX created successfully.")

if __name__ == "__main__":
    desktop_path = r"C:\Users\Ziad.Mahmoud\Desktop\دليل_التجربة_العملية_لمنصة_مركز_أسما.docx"
    local_path = r"c:\Users\Ziad.Mahmoud\Desktop\Asma_Telehealth\دليل_التجربة_العملية_لمنصة_مركز_أسما.docx"
    create_demo_doc(desktop_path)
    create_demo_doc(local_path)
